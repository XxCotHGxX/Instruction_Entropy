from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_mwais_paper():
    doc = Document()
    
    # Set standard AIS margins (approximate)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    # Title
    title = doc.add_heading('THE COMPLEXITY KINK: MAPPING THE FRONTIERS OF AI MARGINAL PRODUCTIVITY VIA INSTRUCTION ENTROPY', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Abstract Section
    doc.add_heading('Abstract', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    abstract_text = (
        "As Large Language Models (LLMs) achieve parity with human benchmarks in discrete tasks, "
        "the structural limits of AI productivity remain poorly quantified. This research introduces two novel "
        "econometric variables: Instruction Entropy (E), representing the ratio of solution information to instruction length, "
        "and Artifact Coupling (kappa), measuring structural coordination complexity. Using the Scale AI Remote Labor Index (RLI) "
        "and 2026 freelance market data, I identify a 'Complexity Kink'---a structural tipping point where AI Marginal Productivity "
        "collapses and human wage premiums accelerate. I apply a Clustered Hedonic Translog model to a decomposed dataset of "
        "156 professional requirements. Preliminary results indicate a statistically significant (p=0.022) non-linear threshold "
        "for coordination complexity, suggesting that human value is concentrated in high-entropy, multi-asset domains. "
        "This framework provides a standardized methodology for monitoring the trajectory of AI-driven labor automation."
    )
    p = doc.add_paragraph(abstract_text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Keywords
    k = doc.add_paragraph()
    k.add_run('Keywords: ').bold = True
    k.add_run('Instruction Entropy, Artifact Coupling, AI Automation, Labor Economics, Complexity Kink.')

    # Introduction
    doc.add_heading('1. Introduction', level=1)
    doc.add_paragraph(
        "The rapid deployment of frontier AI agents in early 2026 has fundamentally disrupted the digital freelance economy. "
        "While LLMs have reached saturation on traditional knowledge-based benchmarks, a significant 'Expert Zone' persists "
        "where human professionals command substantial market premiums. Current literature often attributes this divergence "
        "to subjective 'difficulty.' This study proposes a rigorous alternative: Instruction Entropy (E) and Artifact Coupling (kappa)."
    )
    doc.add_paragraph(
        "Existing measures of AI capability often focus on token length or task duration, metrics which fail to capture the "
        "underlying structural coordination costs inherent in professional work. By defining Instruction Entropy as a measure "
        "of requirements density and Artifact Coupling as a measure of cross-asset dependency, we can mathematically locate the "
        "point where non-biological intelligence hits a productivity ceiling."
    )

    # Methodology
    doc.add_heading('2. Methodology', level=1)
    doc.add_paragraph(
        "To ensure statistical rigor given the small-N nature of high-fidelity professional benchmarks, I utilized the "
        "Scale AI Remote Labor Index (RLI) Public Set. I implemented a three-stage pipeline to isolate the complexity "
        "coordinates of professional labor."
    )

    doc.add_heading('2.1 Metric Definitions', level=2)
    doc.add_paragraph(
        "Instruction Entropy (E): Defined as a Boilerplate-Agnostic MDL (Minimum Description Length) ratio. To isolate "
        "pure inference from template noise, I utilize a mask that strips standardized LaTeX preambles and code imports before tokenization."
    )
    doc.add_paragraph(
        "Artifact Coupling (kappa): Measures coordination complexity across the deliverable structure. It is calculated as a weighted "
        "Structural Complexity Index combining file fan-out and maximum hierarchy depth."
    )

    doc.add_heading('2.2 Requirement Decomposition', level=2)
    doc.add_paragraph(
        "I decomposed 10 foundational RLI projects into 156 discrete professional requirements. This allowed for a shift in "
        "the unit of analysis from 'The Project' to 'The Requirement,' providing the degrees of freedom necessary for high-confidence "
        "modeling while preserving the depth of the original data."
    )

    doc.add_heading('2.3 Econometric Specification', level=2)
    doc.add_paragraph(
        "I estimate labor value using a Clustered Hedonic Translog model. Standard errors are clustered at the project level "
        "to account for intra-task correlation. Wages are anchored to O*NET SOC-specific baseline wages to mitigate self-reporting "
        "bias in freelance completion times."
    )

    # Results
    doc.add_heading('3. Results', level=1)
    doc.add_paragraph(
        "The results identify a statistically significant non-linear threshold for Artifact Coupling. As coordination complexity increases, "
        "the marginal productivity of AI labor enters a sharp decline, leading to the 'Complexity Kink' observed in market pricing data."
    )

    # Table 1
    doc.add_paragraph('Table 1: Clustered Hedonic Translog Coefficients (N=156)', style='Caption')
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Variable'
    hdr_cells[1].text = 'Coefficient'
    hdr_cells[2].text = 'Std. Error'
    hdr_cells[3].text = 'P-Value'
    
    data = [
        ('Intercept', '3.086', '0.152', '0.000'),
        ('log(E)', '0.291', '0.026', '0.000'),
        ('log(kappa)', '-0.011', '0.100', '0.909'),
        ('log(kappa)^2', '-0.139', '0.061', '0.022'),
        ('AI Exposure', '29.955', '5.981', '0.000'),
    ]
    
    for var, coef, se, p in data:
        row_cells = table.add_row().cells
        row_cells[0].text = var
        row_cells[1].text = coef
        row_cells[2].text = se
        row_cells[3].text = p

    # Figure placeholder
    doc.add_paragraph('\n[FIGURE 1: THE COMPLEXITY FRONTIER GRADIENT (KDE)]', style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Figure 1: Concentration of human expert labor in high-entropy domains.', style='Caption').alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Discussion
    doc.add_heading('4. Discussion', level=1)
    doc.add_paragraph(
        "The 'AI Labor Floor' is not a static line but a dynamic frontier. While agentic loops are successfully lowering the cost of execution, "
        "the coordination cost of high-entropy tasks remains a bottleneck for AI. Human value in 2026 is concentrated in the 'Entropy Tail' where "
        "structural orchestration is required."
    )

    # Limitations
    doc.add_heading('5. Limitations', level=1)
    doc.add_paragraph(
        "This study faces selection bias inherent in the RLI public set. Future research must expand the E-kappa framework to wider "
        "freelance datasets to verify coordinate stability. Furthermore, as model reasoning improves, the coordinates of the Kink are expected to shift rightward."
    )

    # References
    doc.add_heading('References', level=1)
    doc.add_paragraph("Mazeika, M., Gatti, A., et al. (2025). 'Remote Labor Index: Measuring AI Automation of Remote Work.' arXiv:2510.26787.", style='List Bullet')
    doc.add_paragraph("Eloundou, T., et al. (2023). 'GPTs are GPTs: An Early Look at the Labor Market Impact Potential of Large Language Models.' OpenAI.", style='List Bullet')
    doc.add_paragraph("National Center for O*NET Development. (2025). 'O*NET 30.0 Database.' US Department of Labor.", style='List Bullet')

    doc.save('D:/ProgD/Instruction Entropy/paper/MWAIS_2026_SUBMISSION.docx')
    print("MWAIS 2026 Submission Draft created as .docx")

if __name__ == "__main__":
    create_mwais_paper()
